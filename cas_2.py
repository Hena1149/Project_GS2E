from docx import Document  # Ajoutez cette ligne
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import random
import spacy
from spacy.lang.fr.stop_words import STOP_WORDS
import streamlit as st
import re
import string
from io import BytesIO
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import pdfminer.high_level
import docx

# Configuration de l'application
st.set_page_config(page_title="Générateur de Cas de test à partir du CDC", layout="wide", page_icon="📑")

# ----------------------------
# FONCTIONS UTILITAIRES
# ----------------------------


# @st.cache_resource
# def load_nlp_model():
#     """Charge le modèle spaCy pour le traitement NLP"""
#     try:
#         nlp = spacy.load("fr_core_news_md")
#         return nlp
#     except Exception as e:
#         st.error(f"Erreur de chargement du modèle NLP: {str(e)}")
#         st.info("Veuillez installer le modèle français avec: python -m spacy download fr_core_news_md")
#         return None
@st.cache_resource
def load_nlp_model():
    """Charge le modèle spaCy pour le traitement NLP"""
    try:
        # Essaye de charger le modèle normalement
        nlp = spacy.load("fr_core_news_md")
        st.success("Modèle NLP chargé avec succès !")
        return nlp
    except OSError:
        try:
            # Si le modèle n'est pas trouvé, propose l'installation
            st.error("Modèle français non trouvé. Installation en cours...")
            import os
            os.system("python -m spacy download fr_core_news_md")
            nlp = spacy.load("fr_core_news_md")
            return nlp
        except Exception as e:
            st.error(f"Échec du chargement : {str(e)}")
            return None
    except Exception as e:
        st.error(f"Erreur inattendue : {str(e)}")
        return None


def extract_business_rules(text, nlp_model):
    """
    Version améliorée de l'extraction des règles de gestion
    Combine regex avancés et analyse NLP avec spaCy
    """
    # 1. Prétraitement du texte
    text = fix_incomplete_lines(text)
    
    # 2. Extraction avec motifs regex améliorés
    patterns = [
        # Structures conditionnelles
        r"(Si|Lorsque|Lorsqu'|Quand|Dès que|En cas de|Au cas où)\b.*?"
        r"(alors|doit|devra|est tenu de|nécessite|implique|entraîne|peut|sera)\b.*?[.;]",
        
        # Obligations
        r"(L'?utilisateur|Le client|Le système|L'?application|Un administrateur)\b.*?"
        r"(doit|est tenu de|devra|a l'obligation de|est responsable de)\b.*?[.;]",
        
        # Interdictions
        r"(Il est interdit|Ne doit pas|Est prohibé|N'est pas autorisé)\b.*?[.;]",
        
        # Contrôles/Validations
        r"(Vérifier|Valider|Contrôler|S'assurer que|Garantir)\b.*?[.;]",
        
        # Conséquences
        r"(En cas de|Si non respect|En cas de non-conformité)\b.*?"
        r"(entraîne|provoque|conduit à|aura pour effet)\b.*?[.;]",
        
        # Droits/autorisations
        r"(L'?utilisateur|Le client|Le prestataire|L'?agent|Le système)\b.*?"
        r"(est autorisé à|peut|a le droit de)\b.*?[.;]"
    ]
    
    rules = set()
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            rule = clean_business_rule(match.group())
            if is_valid_rule(rule):
                rules.add(rule)
    
    # 3. Extraction NLP si le modèle est disponible
    if nlp_model:
        doc = nlp_model(text)
        
        for sent in doc.sents:
            sent_text = sent.text.strip()
            
            # Détection des phrases contenant des marqueurs de règles
            if is_business_rule_sentence(sent):
                rule = clean_business_rule(sent_text)
                if is_valid_rule(rule):
                    rules.add(rule)
    
    # 4. Post-traitement des règles
    rules = clean_short_rules(rules)
    rules = sorted(rules, key=lambda x: len(x), reverse=True)
    
    return rules

# Fonctions utilitaires améliorées
def fix_incomplete_lines(text):
    """Corrige les phrases de conditions coupées en fusionnant les lignes incomplètes"""
    lines = text.split("\n")
    fixed_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.lower().startswith(("si ", "lorsqu'", "quand ", "dès que ", "en cas de ")):
            # Fusion avec la ligne suivante si la phrase n'est pas terminée
            while i + 1 < len(lines) and not line.endswith(('.', ';')):
                i += 1
                next_line = lines[i].strip()
                if next_line:  # Évite d'ajouter des lignes vides
                    line += " " + next_line
        fixed_lines.append(line)
        i += 1
    
    return "\n".join(fixed_lines)

def clean_business_rule(rule_text):
    """Nettoie et formate une règle de gestion"""
    # Suppression des espaces multiples et caractères spéciaux indésirables
    rule_text = re.sub(r"[^\w\sàâäéèêëîïôöùûüç,;.']", " ", rule_text, flags=re.IGNORECASE)
    rule_text = re.sub(r"\s+", " ", rule_text).strip()
    
    # Capitalisation et ponctuation
    if rule_text:
        rule_text = rule_text[0].upper() + rule_text[1:]
        if not rule_text.endswith(('.', ';')):
            rule_text += '.'
    
    return rule_text

def is_valid_rule(rule_text):
    """Valide qu'une règle extraite est complète"""
    words = rule_text.split()
    return (len(words) > 4 and  # Règles d'au moins 5 mots
            not rule_text.startswith(('Comment ', 'Pourquoi ', 'Quand ', 'Où ')) and
            not any(word in rule_text.lower() for word in ['exemple', 'note', 'remarque']))

def is_business_rule_sentence(sent):
    """Détermine si une phrase est une règle métier valide avec spaCy"""
    # Liste des termes déclencheurs
    rule_keywords = {
        'VERB': ['devoir', 'falloir', 'pouvoir', 'interdire', 'autoriser'],
        'NOUN': ['obligation', 'interdiction', 'condition', 'requis', 'validation'],
        'ADJ': ['obligatoire', 'interdit', 'autorisé', 'requis']
    }
    
    # Vérifie la présence de termes clés
    has_keyword = False
    for token in sent:
        if token.text.lower() in ['si', 'alors', 'doit', 'nécessite']:
            has_keyword = True
        if token.pos_ in rule_keywords and token.lemma_ in rule_keywords[token.pos_]:
            has_keyword = True
    
    # Vérifie la structure de la phrase
    has_conditional = any(token.dep_ == "mark" for token in sent)  # Marqueurs de condition
    has_modal = any(token.dep_ == "aux" for token in sent)         # Verbes modaux
    
    return (len(sent) > 5 and 
            (has_keyword or has_conditional or has_modal) and
            not any(ent.label_ == "DATE" for ent in sent.ents))    # Exclut les dates

def clean_short_rules(rules):
    """Filtre les règles trop courtes ou incomplètes"""
    return [rule for rule in rules 
            if len(rule.split()) > 4 and 
            not rule.lower().startswith(('page ', 'article ', 'paragraphe '))]




def clean_rule(rule_text):
    """Nettoie et formate une règle de gestion"""
    rule_text = re.sub(r"\s+", " ", rule_text).strip()
    if not rule_text.endswith('.'):
        rule_text += '.'
    return rule_text


def create_pdc_document(pdc_list):
    """Crée un document Word à partir des PDC"""
    from docx import Document  # Solution alternative si vous ne pouvez pas ajouter l'import global
    doc = Document()
    doc.add_heading('Points de Contrôle (PDC)', level=1)
    for i, pdc in enumerate(pdc_list, 1):
        p = doc.add_paragraph(style='ListBullet')
        p.add_run(f"{i}. {pdc}").bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def extract_text(uploaded_file):
    """Extrait le texte depuis PDF ou DOCX"""
    try:
        file_bytes = uploaded_file.getvalue()
        
        if uploaded_file.type == "application/pdf":
            with BytesIO(file_bytes) as f:
                text = pdfminer.high_level.extract_text(f)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with BytesIO(file_bytes) as f:
                doc = docx.Document(f)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise ValueError("Format non supporté")
            
        return text if text and text.strip() else None
        
    except Exception as e:
        st.error(f"Erreur d'extraction : {str(e)}")
        return None


# Déplacer cette fonction ici (au même niveau que les autres fonctions)
def create_rules_document(rules):
    """Crée un document Word des règles"""
    doc = docx.Document()
    doc.add_heading('Règles de Gestion Identifiées', level=1)
    
    for i, rule in enumerate(rules, 1):
        doc.add_paragraph(f"{i}. {rule}", style='ListBullet')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# def clean_text(text):
#     """Nettoyage basique du texte"""
#     text = text.lower()
#     text = re.sub(r"[^\w\s]", " ", text)  # Supprime la ponctuation
#     text = re.sub(r"\s+", " ", text)      # Espaces multiples -> simple
#     return text.strip()

def clean_text(text, nlp_model, min_word_length=3):
    """
    Nettoyage approfondi du texte avec :
    - Suppression des stopwords
    - Lemmatisation
    - Filtrage par catégorie grammaticale
    - Suppression des mots trop courts
    """
    # Initialisation
    if not text or not nlp_model:
        return ""
    
    # Nettoyage de base
    text = text.lower()
    text = re.sub(r"[^\w\sàâäéèêëîïôöùûüç]", " ", text)  # Garde les caractères accentués
    text = re.sub(r"\s+", " ", text).strip()
    
    # Traitement NLP
    doc = nlp_model(text)
    cleaned_tokens = []
    
    for token in doc:
        # Conditions de filtrage
        if (token.is_stop or 
            token.is_punct or 
            len(token.text) < min_word_length or
            token.pos_ in ["DET", "ADP", "CCONJ", "PRON", "PART"]):
            continue
            
        # Lemmatisation (forme de base)
        lemma = token.lemma_.strip()
        if lemma:
            cleaned_tokens.append(lemma)
    
    return " ".join(cleaned_tokens)

def calculate_frequencies(text):
    """Calcule les fréquences des mots"""
    words = [word for word in text.split() if len(word) > 2]  # Filtre mots courts
    return pd.Series(words).value_counts()


def generate_wordcloud(freq_dict, width=800, height=400, background_color="white", colormap="viridis"):
    """Génère un nuage de mots"""
    fig, ax = plt.subplots(figsize=(10, 5))
    wc = WordCloud(
        width=width,
        height=height,
        background_color=background_color,
        colormap=colormap,
        max_words=100
    ).generate_from_frequencies(freq_dict)
    
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    return fig

# ----------------------------
# NOUVELLES FONCTIONS POUR L'ONGLET PDC
# ----------------------------

def extract_pdc_from_text(text):
    """Extrait les exigences PDC d'un texte"""
    patterns = [
        r"(Vérifier|S['’]assurer|Contrôler|Vérification|Point de contrôle)\b.*?[\.;]",
        r"(Le système doit|Il faut|Il est nécessaire de).*?(vérifier|contrôler|s'assurer)"
    ]
    pdc_list = set()
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            pdc = match.group().strip()
            if len(pdc.split()) > 3:  # Filtre les phrases trop courtes
                if not pdc.endswith('.'):
                    pdc += '.'
                pdc_list.add(pdc)
    return sorted(pdc_list, key=lambda x: len(x), reverse=True)

# Dans votre fonction main() ou au début du script :
if 'nlp' not in st.session_state:
    load_nlp_model()
def generate_pdc_from_rule(rule):
    """Génère un PDC à partir d'une règle de gestion"""
    if 'nlp' not in st.session_state:
        st.error("Modèle NLP non chargé")
        return f"Vérifier que {rule}"
    
    doc = st.session_state.nlp(rule)
    verbs = [token.text for token in doc if token.pos_ == "VERB"]
    action = verbs[0] if verbs else "vérifier"
    return f"{action.capitalize()} que {rule}"


def compare_rules_pdc(rules, pdc_list):
    """Compare les règles avec les PDC existants"""
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(rules + pdc_list)
    similarity = cosine_similarity(tfidf_matrix[:len(rules)], tfidf_matrix[len(rules):])
    return similarity


def create_test_case(pdc, index, is_manual=False):
    """Crée un cas de test à partir d'un PDC"""
    templates = [
        f"Le système doit satisfaire : {pdc}",
        f"Confirmer que {pdc}",
        f"Tester la conformité de : {pdc}"
    ]
    return {
        "ID": f"CT-{index:03d}",
        "Type": "Manuel" if is_manual else "Auto-généré",
        "PDC": pdc,
        "Description": random.choice(templates) if not is_manual else pdc,
        "Étapes": f"1. Préparer l'environnement\n2. Exécuter: {pdc}\n3. Vérifier le résultat",
        "Résultat attendu": f"{pdc} est correctement implémenté"
    }


def create_pdc_document(pdc_list):
    """Crée un document Word à partir des PDC"""
    doc = Document()
    doc.add_heading('Points de Contrôle (PDC)', level=1)
    for i, pdc in enumerate(pdc_list, 1):
        p = doc.add_paragraph(style='ListBullet')
        p.add_run(f"{i}. {pdc}").bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------------------------
# INTERFACE UTILISATEUR
# ----------------------------
st.title("Générateur de Cas de Test")
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📤 Extraction", "🔍 Analyse", "☁️ WordCloud", "📜 Règles", "✅ PDC & Tests"])

with tab1:
    st.header("Extraction de Texte")
    uploaded_file = st.file_uploader("Téléversez un document (PDF ou DOCX)", type=["pdf", "docx"])
    
    if uploaded_file and st.button("Extraire le texte"):
        with st.spinner("Extraction en cours..."):
            extracted_text = extract_text(uploaded_file)
            
            if extracted_text:
                st.session_state.text = extracted_text
                st.success("Texte extrait avec succès !")
                
                with st.expander("Aperçu du texte"):
                    st.text(extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else ""))

# with tab2:
#     st.header("Analyse Textuelle")
    
#     if 'text' not in st.session_state:
#         st.warning("Veuillez d'abord extraire un texte dans l'onglet 'Extraction'")
#     else:
#         with st.spinner("Nettoyage du texte..."):
#             st.session_state.text_clean = clean_text(st.session_state.text)
#             st.session_state.freq = calculate_frequencies(st.session_state.text_clean)
        
#         st.subheader("Fréquence des mots")
#         top_n = st.slider("Nombre de mots à afficher", 5, 50, 20)
#         st.dataframe(st.session_state.freq.head(top_n))

with tab2:
    if 'text' not in st.session_state:
        st.warning("Veuillez d'abord extraire un texte dans l'onglet 'Extraction'")
    else:
        nlp_model = load_nlp_model()  # Charger le modèle
        if not nlp_model:
            st.error("Modèle NLP non disponible pour le nettoyage")
        else:
            with st.spinner("Nettoyage approfondi en cours..."):
                st.session_state.text_clean = clean_text(st.session_state.text, nlp_model)
                st.session_state.freq = calculate_frequencies(st.session_state.text_clean)
            
            st.subheader("Fréquence des mots (nettoyés)")
            top_n = st.slider("Nombre de mots à afficher", 5, 50, 20)
            st.dataframe(st.session_state.freq.head(top_n))

with tab3:
    st.header("Visualisation WordCloud")
    
    if 'text_clean' not in st.session_state:
        st.warning("Veuillez d'abord analyser un texte dans l'onglet 'Analyse'")
    else:
        with st.expander("Paramètres avancés"):
            col1, col2 = st.columns(2)
            with col1:
                width = st.slider("Largeur", 400, 1200, 800, key="wc_width")
                height = st.slider("Hauteur", 200, 800, 400, key="wc_height")
            with col2:
                bg_color = st.color_picker("Couleur de fond", "#FFFFFF", key="wc_bg")
                colormap = st.selectbox("Palette", ["viridis", "plasma", "inferno", "magma", "cividis"], key="wc_cmap")
        
        if st.button("Générer le WordCloud"):
            freq_dict = st.session_state.freq.to_dict()
            fig = generate_wordcloud(
                freq_dict,
                width=width,
                height=height,
                background_color=bg_color,
                colormap=colormap
            )
            
            st.pyplot(fig)
            
            # Téléchargement
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', bbox_inches='tight')
            st.download_button(
                label="💾 Télécharger l'image",
                data=img_buffer.getvalue(),
                file_name="wordcloud.png",
                mime="image/png"
            )
            
with tab4:
    st.header("Extraction des Règles de Gestion")
    nlp_model = load_nlp_model()
    
    if 'text' not in st.session_state:
        st.warning("Veuillez d'abord extraire un texte dans l'onglet 'Extraction'")
    elif not nlp_model:
        st.error("Le traitement NLP n'est pas disponible")
    else:
        if st.button("Extraire les règles", type="primary"):
            with st.spinner("Analyse en cours (cela peut prendre quelques minutes)..."):
                rules = extract_business_rules(st.session_state.text, nlp_model)
                
                if rules:
                    st.session_state.rules = rules
                    st.success(f"{len(rules)} règles identifiées !")
                    
                    # Affichage paginé
                    st.subheader("Règles extraites")
                    items_per_page = 5
                    total_pages = (len(rules) + items_per_page - 1) // items_per_page
                    
                    page = st.number_input("Page", 1, total_pages, 1, 
                                         help="Naviguez entre les pages de résultats")
                    
                    start_idx = (page - 1) * items_per_page
                    end_idx = min(start_idx + items_per_page, len(rules))
                    
                    for i in range(start_idx, end_idx):
                        st.markdown(f"**Règle {i+1}**")
                        st.info(rules[i])
                    
                    # Export des résultats
                    st.subheader("Export des résultats")
                    docx_file = create_rules_document(rules)
                    st.download_button(
                        "📄 Télécharger au format Word",
                        data=docx_file,
                        file_name="regles_gestion.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Option d'analyse supplémentaire
                    with st.expander("Analyse avancée"):
                        st.metric("Nombre total de règles", len(rules))
                        avg_length = sum(len(rule.split()) for rule in rules) / len(rules)
                        st.metric("Longueur moyenne des règles", f"{avg_length:.1f} mots")
                else:
                    st.warning("Aucune règle de gestion n'a été identifiée dans le document")


with tab5:
    st.header("Gestion des Points de Contrôle et Cas de Test")
    nlp = spacy.load("fr_core_news_md")
    
    if 'rules' not in st.session_state:
        st.warning("Veuillez d'abord extraire les règles dans l'onglet 'Règles'")
    else:
        # Section 1: Chargement des PDC existants
        st.subheader("1. Chargement des PDC existants")
        has_pdc = st.radio("Avez-vous des PDC existants à importer ?", 
                          ("Oui, j'ai des PDC existants", "Non, générer des PDC automatiquement"),
                          index=0)
        
        pdc_file = None
        pdc_text = ""
        
        if has_pdc.startswith("Oui"):
            pdc_file = st.file_uploader("Téléversez votre fichier PDC (PDF/DOCX/TXT)", 
                                       type=["pdf", "docx", "txt"], 
                                       key="pdc_uploader")
            
            if pdc_file:
                with st.spinner("Extraction des PDC en cours..."):
                    pdc_text = extract_text(pdc_file)
                    st.session_state.pdc_list = extract_pdc_from_text(pdc_text)
                    
                    if st.session_state.pdc_list:
                        st.success(f"{len(st.session_state.pdc_list)} PDC extraits !")
                        with st.expander("Aperçu des PDC"):
                            for i, pdc in enumerate(st.session_state.pdc_list[:5], 1):
                                st.markdown(f"{i}. {pdc}")
                    else:
                        st.warning("Aucun PDC détecté dans le document")
                        st.session_state.pdc_list = []
        
        # Section 2: Génération des PDC
        st.subheader("2. Génération des PDC")
        nlp_model = load_nlp_model()
        if has_pdc.startswith("Non") or (has_pdc.startswith("Oui") and pdc_file):
            if st.button("Générer/Compléter les PDC", type="primary"):
                with st.spinner("Création des PDC..."):
                    # Initialisation de la liste PDC
                    if 'pdc_list' not in st.session_state:
                        st.session_state.pdc_list = []
                    
                    # Pour les règles sans PDC correspondant
                    if has_pdc.startswith("Oui") and pdc_file:
                        similarity = compare_rules_pdc(st.session_state.rules, st.session_state.pdc_list)
                        threshold = st.slider("Seuil de similarité pour les correspondances", 0.1, 1.0, 0.6)
                        
                        for i, rule in enumerate(st.session_state.rules):
                            if similarity[i].max() < threshold:
                                generated_pdc = generate_pdc_from_rule(rule)
                                st.session_state.pdc_list.append(generated_pdc)
                    else:
                        # Génération automatique complète
                        nlp_model = load_nlp_model()
                        st.session_state.pdc_list = [generate_pdc_from_rule(rule) for rule in st.session_state.rules]
                    
                    st.success(f"{len(st.session_state.pdc_list)} PDC prêts !")
        
        # Section 3: Visualisation et Export
        if 'pdc_list' in st.session_state and st.session_state.pdc_list:
            st.subheader("3. Points de Contrôle")
            
            # Affichage paginé
            pdc_per_page = 5
            total_pages = (len(st.session_state.pdc_list) + pdc_per_page - 1) // pdc_per_page
            page = st.number_input("Page", 1, total_pages, 1)
            
            start_idx = (page - 1) * pdc_per_page
            end_idx = min(start_idx + pdc_per_page, len(st.session_state.pdc_list))
            
            for i in range(start_idx, end_idx):
                st.markdown(f"**PDC {i+1}**")
                st.info(st.session_state.pdc_list[i])
            
            # Export PDC
            st.download_button(
                "📥 Télécharger les PDC (DOCX)",
                data=create_pdc_document(st.session_state.pdc_list),
                file_name="points_de_controle.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Section 4: Génération des Cas de Test
            st.subheader("4. Cas de Test Associés")
            
            if st.button("Générer les Cas de Test"):
                with st.spinner("Création des cas de test..."):
                    st.session_state.test_cases = []
                    
                    for i, pdc in enumerate(st.session_state.pdc_list, 1):
                        is_manual = has_pdc.startswith("Oui") and i <= len(st.session_state.pdc_list)
                        st.session_state.test_cases.append(create_test_case(pdc, i, is_manual))
                    
                    st.success(f"{len(st.session_state.test_cases)} cas de test générés !")
            
            # Affichage des Cas de Test
            if 'test_cases' in st.session_state:
                df_test_cases = pd.DataFrame(st.session_state.test_cases)
                st.dataframe(df_test_cases[["ID", "Type", "PDC", "Description"]])
                
                # Export des Cas de Test
                test_cases_doc = Document()
                test_cases_doc.add_heading('Cas de Test', level=1)
                
                table = test_cases_doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                headers = ["ID", "Type", "PDC", "Description", "Étapes"]
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header
                
                for case in st.session_state.test_cases:
                    row = table.add_row().cells
                    row[0].text = case["ID"]
                    row[1].text = case["Type"]
                    row[2].text = case["PDC"]
                    row[3].text = case["Description"]
                    row[4].text = case["Étapes"]
                
                buffer = BytesIO()
                test_cases_doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    "📥 Télécharger les Cas de Test (DOCX)",
                    data=buffer,
                    file_name="cas_de_test.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ----------------------------
# PIED DE PAGE
# ----------------------------
st.markdown("---")
st.caption("Application développée avec Streamlit - Mise à jour : %s" % pd.Timestamp.now().strftime("%d/%m/%Y"))
